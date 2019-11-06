# from django.core.urlresolvers import reverse  <- for django 1.11
from django.urls import reverse
# for working with pdfs
import PyPDF2

# for rendering html pages
from django.shortcuts import render

# for working with json
import json

# decarators
from django.contrib.auth.decorators import login_required

# random integers
from random import randint

# similarity finding module
import difflib

# working with time
import datetime

# created forms
from .forms import DataForm,DocumentForm

# for working with urls
from urllib.parse import urlencode, urlparse, parse_qs

# regular expressions
import re

# for storing data in system
from django.core.files.storage import FileSystemStorage

# to make html pretty
from lxml.html import fromstring

# accessing webpages
from requests import get
import requests

# extracting html
from bs4 import BeautifulSoup

# io operations
import io
from os.path import splitext
import random

# math functions
import math

# no idea
from collections import Counter

# redirecting webpages
from django.http import HttpResponseRedirect

# models
from .models import Document,Data,randomid,Apidocument,Credits,Subuser,User,Institution

# view classes
from django.views import View
from rest_framework.views import APIView

# django-rest-framework
from rest_framework.parsers import MultiPartParser, FormParser,FileUploadParser
from rest_framework.response import Response
from rest_framework import status
from .serializers import DataSerializer
from rest_framework import viewsets
from io import BytesIO
from django.http import HttpResponse
from django.template.loader import get_template
import xhtml2pdf.pisa as pisa
import docx
from django.http import FileResponse, Http404

# for highlighting pdf
import fitz

from rest_framework.decorators import api_view
MAX_RETRIES = 20
session = requests.Session()
adapter = requests.adapters.HTTPAdapter(max_retries=MAX_RETRIES)
session.mount('https://', adapter)
session.mount('http://', adapter)
#
class ApiViewSet(viewsets.ModelViewSet):
    queryset = Apidocument.objects.all().order_by('-uploaded_at')
    serializer_class = DataSerializer
    # def list(self, request, *args, **kwargs):
    #     return Response({'something': 'my custom JSON'})


class FileView(APIView):
  parser_classes = (MultiPartParser, FormParser)
  def post(self, request, *args, **kwargs):
    file_serializer = DataSerializer(data=request.data)
    if file_serializer.is_valid():
      file_serializer.save()
      return Response(file_serializer.data, status=status.HTTP_201_CREATED)
    else:
      return Response(file_serializer.errors, status=status.HTTP_400_BAD_REQUEST)

def addusers(request):
    curr_user = request.user
    documents = Document.objects.filter(user=request.user)
    if(curr_user.credits.credits<=1000):
        msg = "Sorry you do not have enough credits to add user"
        return render(request,'dashboard1.html',{'msg':msg,'documents':documents,'cli':curr_user})
    if request.method == 'POST' and 'insregisterbutton' in request.POST:
        print("inside elif")
        try:
            user = User.objects.get(username__iexact=request.POST['username'])
            msg = 'User name already exists try different one'
            return render(request,'addusers.html',{'msg':msg})
        except User.DoesNotExist:
            curr_user.credits.credits -= 1000
            curr_user.save()
            user = User.objects.create_user(username=request.POST['username'],password=request.POST['password1'],email=request.POST['email'])
            user.credits.credits = 1000
            user.institution.flag = 0
            user.save()
            # user = User.objects.get(pk = user.id)
            print("inside institution")
            subuserobj = Subuser(user = user,created_by=request.user)
            subuserobj.save()
            subusers = Subuser.objects.filter(created_by = request.user)
            
            return render(request,'subusers.html',{'subusers':subusers})
    return render(request,'addusers.html',{})

def subusers(request):
    subuserlist = Subuser.objects.filter(created_by = request.user)
    return render(request,'subusers.html',{'subusers':subuserlist})


browsers = [
            'Mozilla/5.0 (Windows; U; Windows NT 5.1; en-US; rv:1.9.0.6) Gecko/2009011913 Firefox/3.0.6',
            'Mozilla/5.0 (Macintosh; U; Intel Mac OS X 10.5; en-US; rv:1.9.0.6) Gecko/2009011912 Firefox/3.0.6',
            'Mozilla/5.0 (Windows; U; Windows NT 5.1; en-US; rv:1.9.0.6) Gecko/2009011913 Firefox/3.0.6 (.NET CLR 3.5.30729)',
            'Mozilla/5.0 (X11; U; Linux i686; en-US; rv:1.9.0.6) Gecko/2009020911 Ubuntu/8.10 (intrepid) Firefox/3.0.6',
            'Mozilla/5.0 (Windows; U; Windows NT 6.0; en-US; rv:1.9.0.6) Gecko/2009011913 Firefox/3.0.6',
            'Mozilla/5.0 (Windows; U; Windows NT 6.0; en-US; rv:1.9.0.6) Gecko/2009011913 Firefox/3.0.6 (.NET CLR 3.5.30729)',
            'Mozilla/5.0 (Windows; U; Windows NT 5.1; en-US) AppleWebKit/525.19 (KHTML, like Gecko) Chrome/1.0.154.48 Safari/525.19',
            'Mozilla/4.0 (compatible; MSIE 7.0; Windows NT 5.1; .NET CLR 1.1.4322; .NET CLR 2.0.50727; .NET CLR 3.0.04506.30; .NET CLR 3.0.04506.648)',
            'Mozilla/5.0 (X11; U; Linux x86_64; en-US; rv:1.9.0.6) Gecko/2009020911 Ubuntu/8.10 (intrepid) Firefox/3.0.6',
            'Mozilla/5.0 (X11; U; Linux i686; en-US; rv:1.9.0.5) Gecko/2008121621 Ubuntu/8.04 (hardy) Firefox/3.0.5',
            'Mozilla/5.0 (Macintosh; U; Intel Mac OS X 10_5_6; en-us) AppleWebKit/525.27.1 (KHTML, like Gecko) Version/3.2.1 Safari/525.27.1',
            'Mozilla/4.0 (compatible; MSIE 6.0; Windows NT 5.1; SV1; .NET CLR 1.1.4322)',
            'Mozilla/4.0 (compatible; MSIE 6.0; Windows NT 5.1; SV1; .NET CLR 2.0.50727)',
            'Mozilla/4.0 (compatible; MSIE 7.0; Windows NT 5.1)'
            'Mozilla/5.0 (X11; Ubuntu; Linux x86_64; rv:57.0) Gecko/20100101 Firefox/57.0',
            ]
WORD = re.compile(r'\w+')
headers={
	'User-Agent': browsers[random.randint(0, len(browsers) - 1)],
	'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
	'Accept-Language': 'en-us,en;q=0.5'
	}

	
def get_cosine(vec1, vec2):
     intersection = set(vec1.keys()) & set(vec2.keys())
     numerator = sum([vec1[x] * vec2[x] for x in intersection])

     sum1 = sum([vec1[x]**2 for x in vec1.keys()])
     sum2 = sum([vec2[x]**2 for x in vec2.keys()])
     denominator = math.sqrt(sum1) * math.sqrt(sum2)

     if not denominator:
        return 0.0
     else:
         d1 = float(numerator) / denominator
         print(d1)
         return d1
def text_to_vector(text):
     words = WORD.findall(text)
     return Counter(words)
def calplag(file1_data,file2_data):

    text1 = file1_data
    text2  = file2_data
    vector1 = text_to_vector(text1)
    vector2 = text_to_vector(text2)
    cosine = get_cosine(vector1, vector2)
    print ('Cosine:', cosine)
    return cosine


def dataextract(urllist,textss,x):
    sentenceslist = []
    datalist = []
    similarity_ratio = []
    inptextsent = []
    randomfiles = []


    m = re.split(r'(?<=[^A-Z].[.?]) +(?=[A-Z])', textss)
    for senten in m:
        sentenceslist.append(senten)
    if x!=1:
    	toplist = urllist[:5]
    else:
    	toplist = urllist[:1]
    for l in toplist:
        f = io.open('file1.txt','w', encoding="utf-8")
        f.truncate()
        f.close()
        rad = randint(0, 10000000000000)
        randomfiles.append(rad)

        print(l)
        if (l!='/' and len(l)>2):
            parsed = urlparse(l)
            root, ext = splitext(parsed.path)
            if ext=='.pdf':
                html = get('http://%s'%(l),verify = False,headers = headers,stream=True)
                with open("syed.pdf","wb") as pdf:
                    for chunk in html.iter_content(chunk_size=1024):

                         # writing one chunk at a time to pdf file
                         if chunk:
                             pdf.write(chunk)
                    pdf.close()
                    docdata = open('syed.pdf','rb')
                    pdfReader = PyPDF2.PdfFileReader(docdata)

                    num_pages = pdfReader.numPages
                    count = 0
                    text = ""

                    while count < num_pages:
                        pageObj = pdfReader.getPage(count)
                        count +=1
                        text += pageObj.extractText()

                    z = io.open('file1.txt','a', encoding="utf-8")
                    z.write(text)
                    z.close()
                    z = io.open('media/%d.txt'%(rad),'a', encoding="utf-8")
                    z.write(text)
                    z.close()
                    datalist.append(text)
                    m = re.split(r'(?<=[^A-Z].[.?]) +(?=[A-Z])', text)
                    for senten in m:
                        sentenceslist.append(senten)

                    # print (text)
            else:
                html = get('http://%s'%(l),verify = False,headers = headers).text
                soup = BeautifulSoup(html)
                for script in soup(["script", "style"]):
                    script.extract()
                    # get text
                text = soup.get_text()
                # break into lines and remove leading and trailing space on each
                lines = (line.strip() for line in text.splitlines())
                # break multi-headlines into a line each
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                # drop blank lines
                text = '\n'.join(chunk for chunk in chunks if chunk)
                z = io.open('file1.txt','a', encoding="utf-8")
                z.write(text)
                z.close()
                z = io.open('media/%d.txt'%(rad),'a', encoding="utf-8")
                z.write(text)
                z.close()
                # print(text)
                datalist.append(text)
                m = re.split(r'(?<=[^A-Z].[.?]) +(?=[A-Z])', text)
                for senten in m:
                    sentenceslist.append(senten)
            d1 = io.open('file1.txt','r', encoding="utf-8")

            file1_data = d1.read()
            file2_data = textss
            similarity_ratio.append(calplag(file1_data,file2_data)*100)
    #similarity_ratio = SequenceMatcher(None,file1_data,file2_data).ratio()
            d1.close()
            print("\n\n\n\n")
            print(randomfiles)
            # print(sentenceslist)

    return similarity_ratio,datalist,sentenceslist,randomfiles



def geturls(req,text,x,filename):
    urllist = []
    abstractlist = []
    list3 = []
    count = len(re.findall(r'\w+', text))
    # creditbalance = round(count/100,0)
    # credits = Credits.objects.filter(user = req.user)
    # if(credits < credits-creditsbalanceance):
    #     return render(req,'dashboard1.html',{})

    # credits = credits-creditsbalance
    # creditsobject =  Credits.objects.get(pk = req.user.id)
    # creditsobject.credits = credits
    # creditsobject.save()
    # wordslist = text.split(" ")
    # startindex = 0
    # endindex = 500
    # lenofwordslist = len(wordslist)


    # here we are scraping from google
    # raw = get("https://www.google.com/search?q=%s"%(text),headers = headers,verify=False).text
    # page = fromstring(raw)
    # for result in page.cssselect(".r a"):
    #     url = result.get("href")
    #     if url.startswith("/url?"):
    #         url = parse_qs(urlparse(url).query)['q']
    #     print(url[0])
    #
    #     urllist.append(url[0].strip('https://'))
    # soup = BeautifulSoup(raw, "lxml")
    #
    # for s in soup.findAll('span', {'class': 'st'}):
    #     print(s.text)
    #     abstractlist.append(s.text)111
    #     list3 = list(zip(urllist,abstractlist))
    # result = dataextract(urllist,text)
    # print(list3)
    # print('\n\n\n\nscrapped from google')
    # if not wordslist:
    #     while(startindex<endindex):
    #         raw = get("http://www.ask.com/web?q=%s&o=0&qo=homepageSearchBox"%(" ".join(wordslist[startindex:endindex])),headers = headers,verify = False).text
    #         soup = BeautifulSoup(raw,'lxml')
    #         resultobj = soup.findAll('div', {'class': 'PartialSearchResults-item'})
    #         abstractobj = soup.findAll('p',{'class':'PartialSearchResults-item-abstract'})
    #         for resul in resultobj:
    #             resindex = resul.find('p')
    #             if resindex!=None:
    #                 urllist.append(resindex.text)
    #                 break
    #         for abst in abstractobj:
    #             if abst != None:
    #                 abstractlist.append(abst.text)
    #                 break
    #         startindex+=500
    #         endindex+=500







    if not list3:
        raw = get("http://www.ask.com/web?q=%s&o=0&qo=homepageSearchBox"%(text[:1000]),headers = headers,verify = False).text
        soup = BeautifulSoup(raw,'lxml')
        resultobj = soup.findAll('div', {'class': 'PartialSearchResults-item'})
        abstractobj = soup.findAll('p',{'class':'PartialSearchResults-item-abstract'})
        for resul in resultobj:
            resindex = resul.find('p')
            if resindex!=None:
                urllist.append(resindex.text)
        for abst in abstractobj:
            if abst != None:
                abstractlist.append(abst.text)

        result,datalist,sentence,randomfiles = dataextract(urllist,text,x)


        list3 = list(zip(urllist,abstractlist,result,datalist,randomfiles))
        if(len(result)>0):
        	result.sort()
        	maxterm = round(result[-1],2)
        	# print('dfdf')
        else:
        	maxterm = result
        	#print('dfdf')
        # print(maxterm)
        newdoc = Document.objects.filter(document = filename)
        # newdoc.result = maxterm
        # newdoc.save()
        for d in newdoc:
            d.result = maxterm
            d.save()

        print(randomfiles)
        print('\n\n\n\nscraped from ask')


    return render(req,'display.html',{'links':list3,'result':maxterm,'text':text,'sentences':sentence})

# Create your views hereself.
outputdisplay = ""
class Check(View):
    def get(self, request, *args, **kwargs):
        #bg_image = 'https://upload.wikimedia.org/wikipedia/commons/0/05/20100726_Kalamitsi_Beach_Ionian_Sea_Lefkada_island_Greece.jpg'
        return render(request,'index1.html',{})
        #return render(request, "shortener/home.html", {}) # Try Django 1.8 & 1.9 http://joincfe.com/youtube

    def post(self, request, *args, **kwargs):
        #print(self)
        if request.method == "POST":
            form = DataForm(request.POST)
            if 'quickcheck' in request.POST:
                print('quickcheck')
                if not request.user.is_authenticated:
                    return HttpResponseRedirect(reverse('checker:login'))
                    return None
                if form.is_valid():
                    post = form.save(commit =False)
                    post.user = request.user
                    post.save()
                    cd = form.cleaned_data
                    text1 = cd['data']
                    # text1 = ''.join(e for e in text1 if e.isalnum())
                    text1 = re.sub("\s\s+" , " ", text1)
                    print("\n\n\n cleaned text")
                    print(text1)
                    x = io.open('file2.txt','w', encoding="utf-8")
                    x.truncate()
                    x.close()
                    x = io.open('file2.txt','a', encoding="utf-8")
                    x.write(text1)
                    x.close()
                    x = io.open('sentences.txt','w', encoding="utf-8")
                    x.truncate()
                    x.close()
                    x = io.open('sentences.txt','a', encoding="utf-8")
                    x.write(text1)
                    x.close()

                    return geturls(request,text = text1,x=1,filename=None)
            elif 'deepcheck' in request.POST:
                print('deepcheck')
                if not request.user.is_authenticated:
                    return HttpResponseRedirect(reverse('checker:login'))
                    return None
	        	# form = DataForm(request.POST)
                if form.is_valid():
                    post = form.save(commit =False)
                    post.user = request.user
                    post.save()
                    cd = form.cleaned_data
                    text1 = cd['data']
                    text1 = re.sub("\s\s+" , " ", text1)
                    print("\n\n\n cleaned text")
                    print(text1)
                    x = io.open('file2.txt','w', encoding="utf-8")
                    x.truncate()
                    x.close()
                    x = io.open('file2.txt','a', encoding="utf-8")
                    x.write(text1)
                    x.close()
                    x = io.open('sentences.txt','w', encoding="utf-8")
                    x.truncate()
                    x.close()
                    x = io.open('sentences.txt','a', encoding="utf-8")
                    x.write(text1)
                    x.close()
                    return geturls(request,text = text1,x=0,filename=None)
        else:
            form = DataForm()

        return render(request,'index1.html',{'form':form})

def indexpage(request):
    return render(request,'index3.html',{})

@login_required(login_url='/login')
def dashboard(request):
    if  request.user.is_authenticated:
        userr = request.user
        documents = Document.objects.filter(user=request.user)
    else:
        documents = {}
    return render(request,'dashboard1.html',{'documents':documents,'cli':userr})
def model_form_upload(request):
    global outputdisplay
            
    # print(randomid)
    if request.method == 'POST':
        if not request.user.is_authenticated:
            return HttpResponseRedirect(reverse('checker:login'))
            return None
        myfile = request.FILES['document']
        newdoc = Document(document = request.FILES['document'])
        newdoc.user = request.user
        newdoc.save()
        fs = FileSystemStorage()
        document = Document(
        description='',
        document=myfile.name,
        uploaded_at=datetime.datetime.now(),)
        document.save()
        filename = fs.save(myfile.name, myfile)
        uploaded_file_url = fs.url(filename)
        # print(filename)

        fileval = str(filename).split('.')
        print(fileval)
        outputdisplay = myfile.name
        print(outputdisplay)
        if fileval[-1]=='pdf' or fileval[-1]=='PDF':
            # print(fileval)
            docdata = io.open('media/%s'%(filename),'rb')
            pdfReader = PyPDF2.PdfFileReader(docdata)
            x = io.open('file2.txt','w', encoding="utf-8")
            x.truncate()
            x.close()
            #fileReader.numPages
            num_pages = pdfReader.numPages
            count = 0
            docdata1 = ""

            while count < num_pages:
                pageObj = pdfReader.getPage(count)
                count +=1
                docdata1 += pageObj.extractText()
            docdata1 = re.sub("\s\s+" , " ", docdata1)
            # docdata1 = ''.join(e for e in docdata1 if e.isalnum())
            x = io.open('file2.txt','a', encoding="utf-8")
            x.write(docdata1)
            x.close()
            x = io.open('sentences.txt','w', encoding="utf-8")
            x.truncate()
            x.close()
            x = io.open('sentences.txt','a', encoding="utf-8")
            x.write(docdata1)
            x.close()
            print(docdata1)
            docdata.close()
        elif fileval[-1]=='txt':
            docdata = io.open('media/%s'%(filename),'r')
            docdata1 = docdata.read()
            docdata.close()
            docdata1 = re.sub("\s\s+" , " ", docdata1)
            x = io.open('file2.txt','w', encoding="utf-8")
            x.truncate()
            x.close()
            x = io.open('file2.txt','a', encoding="utf-8")
            x.write(docdata1)
            x.close()
            x = io.open('sentences.txt','w', encoding="utf-8")
            x.truncate()
            x.close()
            x = io.open('sentences.txt','a', encoding="utf-8")
            x.write(docdata1)
            x.close()
            # print(docdata1)
        elif fileval[-1] == 'docx':
            doc = docx.Document('media/%s'%(filename))
            fulltext = []
            for para in doc.paragraphs:
                fulltext.append(para.text)

            docdata1 ='\n'.join(fulltext)
            # docdata1 = ""
            # for i in fulltext:
            #     docdata1+="\n"+str(i)
            x = io.open('file2.txt','w', encoding="utf-8")
            x.truncate()
            x.close()
            x = io.open('file2.txt','a', encoding="utf-8")
            x.write(docdata1)
            x.close()
            x = io.open('sentences.txt','w', encoding="utf-8")
            x.truncate()
            x.close()
            x = io.open('sentences.txt','a', encoding="utf-8")
            x.write(docdata1)
            x.close()




        return geturls(request,text = docdata1,x=0,filename = myfile)


            #return redirect('home')
    else:
        form = DocumentForm()
    return render(request, 'upload.html', {
        'form': form
    })

from difflib import SequenceMatcher
from heapq import nlargest as _nlargest

def get_close_matches_indexes(word, possibilities, n=3, cutoff=0.6):
    """Use SequenceMatcher to return a list of the indexes of the best 
    "good enough" matches. word is a sequence for which close matches 
    are desired (typically a string).
    possibilities is a list of sequences against which to match word
    (typically a list of strings).
    Optional arg n (default 3) is the maximum number of close matches to
    return.  n must be > 0.
    Optional arg cutoff (default 0.6) is a float in [0, 1].  Possibilities
    that don't score at least that similar to word are ignored.
    """

    if not n >  0:
        raise ValueError("n must be > 0: %r" % (n,))
    if not 0.0 <= cutoff <= 1.0:
        raise ValueError("cutoff must be in [0.0, 1.0]: %r" % (cutoff,))
    result = []
    s = SequenceMatcher()
    s.set_seq2(word)
    for idx, x in enumerate(possibilities):
        s.set_seq1(x)
        if s.real_quick_ratio() >= cutoff and \
           s.quick_ratio() >= cutoff and \
           s.ratio() >= cutoff:
            result.append((s.ratio(), idx))

    # Move the best scorers to head of list
    result = _nlargest(n, result)

    # Strip scores for the best n matches
    return [x for score, x in result]

def viewdocument(request,string):
    doctype = string.split('.')
    try:
        if(doctype[-1] == 'pdf'):
            return FileResponse(open('media/%s'%(string), 'rb'), content_type='application/pdf')
        elif (doctype[-1] =='docx'):
            return FileResponse(open('media/%s'%(string), 'rb'), content_type='application/ms-word')
    except FileNotFoundError:
        raise Http404()


def show(request,string):
    global outputdisplay
    # print(output)
    # viewdocument(request,outputdisplay)
    # print()
    sent=[]
    textsent = []
    matchsent = []
    fulllist = []


    print(string)
    d1 = io.open('media/%s.txt'%(string),'r', encoding="utf-8")
    text = d1.read()
    text =  re.sub('[!@#$\n\t0123456789]','',text)
    text = text.replace("'","")
    text = text.replace("(","")
    text = text.replace(")","")
    text = text.replace("-","")
    text = text.replace("_","")
    text = text.replace('"','')
    text = text.replace('[','')
    text = text.replace(']','')
    text = text.replace('^','')
    text = text.replace('?','')
    text = text.replace('*','')
    m = re.split(r'(?<=[^A-Z].[.?]) +(?=[A-Z])', text)
    for senten in m:
        textsent.append(senten)
    d1.close()
    d1 = io.open('sentences.txt','r', encoding="utf-8")
    file1_data = d1.read()
    # file1_data =  re.sub('[!@#$\n\t0123456789-]','',file1_data)
    # file1_data=file1_data.replace("'","")
    # file1_data=file1_data.replace("(","")
    # file1_data=file1_data.replace(")","")
    # file1_data=file1_data.replace("-","")
    # file1_data=file1_data.replace("_","")
    # file1_data=file1_data.replace('"','')
    # file1_data=file1_data.replace('[','')
    # file1_data=file1_data.replace(']','')
    # file1_data=file1_data.replace('^','')
    # file1_data=file1_data.replace('?','')
    m = re.split(r'(?<=[^A-Z].[.?]) +(?=[A-Z])', file1_data)
    for senten in m:
        sent.append(senten)
    d1.close()
    print(len(sent))
    for se in textsent:
        # print(se)
        gg = get_close_matches_indexes(se[:200],m,n=4,cutoff=0.3)
        # gg = " ".join(gg)
        if(len(gg)>0):
            for i in gg:
                if(m[i][0]!='0'):
                    m[i] = '0'+m[i]

        # print(textsent)
        # textsent = [word.replace(gg,'<mark>'+gg+'</mark>') for word in textsent]

        # matchsent.append(gg)

        
        # textsent = map(lambda x: str.replace(x, gg, "<mark>"+ gg +"</mark>"), textsent)
        # matchsent.append(gg)
    # fd = list(zip(textsent,matchsent))
    # for se in matchsent:
    #     for s in se:
    #         fulllist.append(s)
    # # print(fulllist)
    json_list = json.dumps(fulllist)
    print(len(textsent))
    doc  = fitz.open("media/%s"%outputdisplay)
    # r = list(range(len(doc)))   
    for temsentence in m:
        if(temsentence[0] == '0'):
            for page in doc:
                text_instances = page.searchFor(temsentence[1:])
                print(temsentence[1:])
                for inst in text_instances:
                    highlight = page.addHighlightAnnot(inst)

    doc.save("media/%s.pdf"%string,garbage=4,deflate  = True,clean = True)
    # request,'show.html',{'text':text,'sentences':json_list,}
    # doctype = outputdisplay.split('.')
    try:
        return FileResponse(open("media/%s.pdf"%string, 'rb'), content_type='application/pdf')
    except FileNotFoundError:
        raise Http404()






    # return Render.render('show.html',{"text":textsent,"request":request})
